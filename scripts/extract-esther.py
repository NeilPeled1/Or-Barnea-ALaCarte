#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Extract text from docx and xlsx files for Esther Bar Dizengoff project."""

import os
import json

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

DOWNLOADS = "/Users/neil.peled/Downloads"

# Esther Bar files - filename -> key
FILES = [
    ("צ'ק ליסט הכנות.docx", "checklist_prep"),
    ("פיק אפ ערב קפה אסתר.docx", "pickup_evening"),
    ("פיק אפ סלטים אסתר בר .docx", "pickup_salads"),
    ("פיק אפ מחלקת לחם חם ובוקר בריאות.docx", "pickup_bread_health"),
    ("פיק אפ כריכים אסתר בר .docx", "pickup_sandwiches"),
    ("עץ מוצר אסתר מעודכן .xlsx", "product_tree"),
    ("סלטים .docx", "salads"),
    ("מסמך עידכון תפריט קפה אסתר .docx", "menu_update"),
    ("מסמך ללא שם.docx", "unnamed"),
    ("אסתר בר ספר מתכונים.docx", "recipe_book"),
    ("תפריט ערב אסתר לאחר פרזנטציה 31_10.docx", "menu_evening"),
    ("תפריט הרצה מעודכן אסתר .docx", "menu_running"),
    ("קפה סינמה קונספט קולינרי.docx", "cinema_concept"),
    ("מסקנות תפריט.docx", "menu_conclusions"),
    ("כניסה להקמה אסתר בר .docx", "setup_entry"),
    ("טיוטה לתפריט קפה סינמה.docx", "draft_cinema"),
    ("טיוטא לתפריט קפה סינמה.docx", "draft_cinema_alt"),
    ("הזמנות לפתיחה אסתר בר.xlsx", "opening_orders"),
    ("הזמנות לפרזנטצייה ערב .xlsx", "orders_presentation"),
    ("הזמנות לטעימות 9.11.xlsx", "orders_tasting"),
    ("הזמנות אסתר בר .xlsx", "orders_esther"),
    ("Copy of קובץ_תפריט_מעודכן_V1(1).xlsx", "menu_updated_v1"),
]


def extract_docx(path):
    if not Document:
        return "[python-docx not installed]"
    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    except Exception as e:
        return f"[Error: {e}]"


def extract_xlsx(path):
    if not openpyxl:
        return "[openpyxl not installed]"
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            rows.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                vals = [str(v) if v is not None else "" for v in row]
                if any(vals):
                    rows.append("\t".join(vals))
        return "\n".join(rows)
    except Exception as e:
        return f"[Error: {e}]"


def main():
    output = {}
    for filename, key in FILES:
        path = os.path.join(DOWNLOADS, filename)
        if not os.path.exists(path):
            output[key] = {"file": filename, "content": "[File not found]", "path": path}
            continue
        if filename.endswith(".docx"):
            content = extract_docx(path)
        elif filename.endswith(".xlsx"):
            content = extract_xlsx(path)
        else:
            content = "[Unsupported format]"
        output[key] = {"file": filename, "content": content[:15000] if content else ""}
    print(json.dumps(output, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()

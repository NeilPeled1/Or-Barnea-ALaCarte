#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Extract text from docx and xlsx files for Sheffield project."""

import os
import json
import sys

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

DOWNLOADS = "/Users/neil.peled/Downloads"

FILES = [
    ("פיק אפ ראשונות שפילד.docx", "pickup_starters"),
    ("צ'קליסט הכנות לפי פסים שפילד האוס.xlsx", "checklist_stations"),
    ("תפקידים וציפיות ים - מנהל מסעדה.docx", "roles_manager"),
    ("תפריט קינוחים שפילד סופי.docx", "menu_desserts"),
    ("תפריט שפילד סופי .docx", "menu_main"),
    ("פיק אפ פס פיצה שפילד.docx", "pickup_pizza"),
    ("פיק אפ עיקריות שפילד.docx", "pickup_mains"),
    ("ספר מתכונים שפילד 2026.docx", "recipe_book"),
    ("סיכום שפילד 20.1.docx", "summary"),
    ("מסמך שקילות.docx", "equivalence"),
    ("יומן נקיונות שפילד.docx", "cleaning_journal"),
    ("הזמנות שפילד לפרזנטציה.xlsx", "orders_presentation"),
    ("ספר מתכונים שפילד (1).docx", "recipe_book_alt"),
    ("שפילד תפריט עם מחירים .docx", "menu_prices"),
    ("רשימת כלים וכמויות שפילד.xlsx", "tools_quantities"),
    ("עץ מוצר שפילד 2026.xlsx", "product_tree"),
    ("נתוני מכירות - אוקטובר נובמבר דצמבר 2025.xlsx", "sales_data"),
    ("ניהול פרויקט שפילד.docx", "project_management"),
    ("הצעת מחיר שפילד בר .docx", "bar_quote"),
    ("אקסל ספקים - 2026.xlsx", "suppliers"),
    ("הזמנות לפתיחה שפילד.xlsx", "opening_orders"),
]

def extract_docx(path):
    if not Document:
        return "[python-docx not installed]"
    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    except Exception as e:
        return f"[Error: {e}]"

def extract_xlsx(path):
    if not openpyxl:
        return "[openpyxl not installed]"
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            rows.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                vals = [str(v) if v is not None else "" for v in row]
                if any(vals):
                    rows.append("\t".join(vals))
        return "\n".join(rows)
    except Exception as e:
        return f"[Error: {e}]"

def main():
    output = {}
    for filename, key in FILES:
        path = os.path.join(DOWNLOADS, filename)
        if not os.path.exists(path):
            output[key] = {"file": filename, "content": "[File not found]", "path": path}
            continue
        if filename.endswith(".docx"):
            content = extract_docx(path)
        elif filename.endswith(".xlsx"):
            content = extract_xlsx(path)
        else:
            content = "[Unsupported format]"
        output[key] = {"file": filename, "content": content[:15000]}
    print(json.dumps(output, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main()
